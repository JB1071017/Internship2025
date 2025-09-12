from flask import Flask, render_template, request, send_file
from google.generativeai import GenerativeModel, configure
import tempfile
from docx import Document
import os

app = Flask(__name__, template_folder="templates", static_folder="static")

# Configure Gemini API
configure(api_key="AIzaSyBsGOw8urROMGEISaAYV8_omMk-a5JoQ3o")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_policy():
    data = request.json
    
    # Initialize Gemini model
    model = GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    Generate a comprehensive 3-page security policy document for {data['companyName']} in the {data['industry']} industry.
    Policy Scope: {data['policyScope']}
    Compliance Requirements: {data['complianceRequirements'] or "None specified"}
    Special Requirements: {data['specialRequirements'] or "None"}
    
    Format the output with proper headings and sections. Include:
    1. Purpose and Scope
    2. Roles and Responsibilities
    3. Access Control
    4. Data Protection
    5. Incident Response
    6. Compliance
    7. Policy Review
    
    Ensure the document is professional and approximately 3 pages when converted to Word.
    """
    
    response = model.generate_content(prompt)
    return {'policy': response.text}

@app.route('/download')
def download_policy():
    company_name = request.args.get('company', 'SecurityPolicy')
    policy_content = request.args.get('content', '')
    
    # Create Word document
    doc = Document()
    doc.add_heading(f'{company_name} Security Policy', 0)
    
    for line in policy_content.split('\n'):
        if line.strip() == '':
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return send_file(
        temp_file.name,
        as_attachment=True,
        download_name=f'{company_name}_Security_Policy.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

